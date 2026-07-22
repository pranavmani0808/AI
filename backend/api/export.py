from fastapi import APIRouter, HTTPException, status, Query
from fastapi.responses import Response, StreamingResponse
from sqlalchemy import select
from datetime import datetime
import io

from backend.database.postgres import AsyncSessionLocal
from backend.database.models import ResearchSessionModel, GeneratedAnswerModel, CrawlResult

router = APIRouter(prefix="/research", tags=["export"])

def generate_markdown_report(query: str, answer: str, sources: list) -> str:
    md = f"# Research Report: {query}\n\n"
    md += f"*Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}*\n\n"
    md += "## Answer\n\n"
    md += f"{answer}\n\n"
    
    if sources:
        md += "## Sources Analyzed\n\n"
        for idx, src in enumerate(sources):
            title = src.get("title") or "Source Link"
            md += f"- **[{idx+1}]** [{title}]({src.get('url')})\n"
            
    return md

def generate_pdf_report(query: str, answer: str, sources: list) -> bytes:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
    story = []
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontSize=20,
        leading=24,
        textColor=colors.HexColor('#4c1d95'),
        spaceAfter=15
    )
    h2_style = ParagraphStyle(
        'DocH2',
        parent=styles['Heading2'],
        fontSize=14,
        leading=18,
        textColor=colors.HexColor('#1e1b4b'),
        spaceBefore=12,
        spaceAfter=6
    )
    body_style = ParagraphStyle(
        'DocBody',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#27272a'),
        spaceAfter=8
    )
    meta_style = ParagraphStyle(
        'DocMeta',
        parent=styles['Normal'],
        fontSize=9,
        leading=12,
        textColor=colors.HexColor('#71717a'),
        spaceAfter=15
    )
    
    story.append(Paragraph(f"Research Report: {query}", title_style))
    story.append(Paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}", meta_style))
    story.append(Spacer(1, 10))
    
    story.append(Paragraph("Answer", h2_style))
    
    # Simple clean tags mapping for reportlab
    clean_ans = answer.replace("**", "<b>").replace("__", "<i>").replace("\n", "<br/>")
    story.append(Paragraph(clean_ans, body_style))
    story.append(Spacer(1, 15))
    
    if sources:
        story.append(Paragraph("Sources Analyzed", h2_style))
        for idx, src in enumerate(sources):
            bullet = f"<b>[{idx+1}]</b> {src.get('title', 'Source')} - <font color='#4c1d95'><u>{src.get('url', '')}</u></font>"
            story.append(Paragraph(bullet, body_style))
            
    doc.build(story)
    return buffer.getvalue()

def generate_docx_report(query: str, answer: str, sources: list) -> bytes:
    from docx import Document
    buffer = io.BytesIO()
    doc = Document()
    
    doc.add_heading(f"Research Report: {query}", 0)
    doc.add_paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    
    doc.add_heading("Answer", level=1)
    doc.add_paragraph(answer)
    
    if sources:
        doc.add_heading("Sources Analyzed", level=1)
        for idx, src in enumerate(sources):
            doc.add_paragraph(f"[{idx+1}] {src.get('title', 'Source')} - {src.get('url', '')}")
            
    doc.save(buffer)
    return buffer.getvalue()

@router.get("/{id}/export")
async def export_research_session(id: int, format: str = Query("markdown", pattern="^(markdown|pdf|docx)$")):
    """
    Exports a completed research session to Markdown, PDF, or DOCX formats.
    Loads entirely from local Postgres database session metrics.
    """
    async with AsyncSessionLocal() as session:
        rs = await session.get(ResearchSessionModel, id)
        if not rs:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Research session {id} not found."
            )
            
        ans_stmt = select(GeneratedAnswerModel).where(GeneratedAnswerModel.search_id == id)
        ans_res = await session.execute(ans_stmt)
        answer_db = ans_res.scalars().first()
        
        crawl_stmt = select(CrawlResult).where(CrawlResult.search_id == id).order_by(CrawlResult.id)
        crawl_res = await session.execute(crawl_stmt)
        pages = crawl_res.scalars().all()
        
        sources_list = [
            {"title": p.title or "Discovered Source", "url": p.url}
            for p in pages
        ]
        
        answer_text = answer_db.answer if answer_db else "No grounded answer saved for this session."
        
        if format == "markdown":
            md_content = generate_markdown_report(rs.query, answer_text, sources_list)
            return Response(
                content=md_content,
                media_type="text/markdown",
                headers={"Content-Disposition": f"attachment; filename=\"research_{id}.md\""}
            )
        elif format == "pdf":
            try:
                pdf_bytes = generate_pdf_report(rs.query, answer_text, sources_list)
                return Response(
                    content=pdf_bytes,
                    media_type="application/pdf",
                    headers={"Content-Disposition": f"attachment; filename=\"research_{id}.pdf\""}
                )
            except Exception as pdf_err:
                print(f"Failed to compile PDF: {pdf_err}")
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail="Failed to generate PDF document."
                )
        elif format == "docx":
            try:
                docx_bytes = generate_docx_report(rs.query, answer_text, sources_list)
                return Response(
                    content=docx_bytes,
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": f"attachment; filename=\"research_{id}.docx\""}
                )
            except Exception as docx_err:
                print(f"Failed to compile DOCX: {docx_err}")
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail="Failed to generate Word document."
                )
